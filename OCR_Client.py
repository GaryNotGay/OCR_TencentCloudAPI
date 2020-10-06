# @Author  :  lijishi
# @Contact :  lijishi@emails.bjut.edu.cn
# @Software:  Pycharm & Python 3.7
# @EditTime:  Oct 6, 2020
# @Version :  2.0
# @describe:  A Full Client Including Confirm Availability And Text Identify
# @LICENSE :  GNU GENERAL PUBLIC LICENSE Version 3

# References
# https://blog.csdn.net/c_chuxin/article/details/100586079

'''
11:通用印刷体识别 12:通用印刷体识别（高速版）13:通用印刷体识别（高精度版） 14:通用手写体识别 15:英文识别
21:身份证识别 22:营业执照识别 23:银行卡识别 24:名片识别
31:增值税发票识别 32:运单识别
41:驾驶证识别 42:车牌识别 43:车辆VIN码识别 44:行驶证识别
51:算式识别 52:表格识别
'''

import socket
import requests
import time
import sys
import os
import base64
import tkinter as tk
import tkinter.filedialog
from tkinter import messagebox
from tkinter import scrolledtext
from tkinter import ttk
from tkinter import END
from tkinter import *
from PIL import Image, ImageTk
from docx import Document
from docx.oxml.ns import qn
from Tencent_OCR import Tencent_OCR_Summary
from Tencent_OCR import Tencent_OCR_BusinessCard
from Tencent_OCR import Tencent_OCR_VehicleLicense
from Tencent_OCR import Tencent_OCR_IDCard
from Text_Arrange import Text_Arrange_Summary
from Text_Arrange import Text_Arrange_BusinessCard_Picture
from Text_Arrange import Text_Arrange_VehicleLicense_Warning
from Text_Arrange import Text_Arrange_DriverLicense_Warning
from Text_Arrange import Text_Arrange_IDCard_Warning
from Text_Arrange import Text_Arrange_IDCard_IDCardCut
from Text_Arrange import Text_Arrange_IDCard_PortraitCut

icon64en = r'AAABAAEAQEAAAAEAIAAoQgAAFgAAACgAAABAAAAAgAAAAAEAIAAAAAAAAEAAAAAAAAAAAAAAAAAAAAAAAAD/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////4uL4/8rK8v/KyvL/5ub5/////////////////////////////////9fX9f/BwfD////////////////////////////////////////////////////////////////////////////19fz/vr7v///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////ExPH/Pz/R/ysrzP8rK8z/KyvM/ysrzP9RUdX/1dX1/////////////////7m57v81Nc7/NTXO/9ra9v//////////////////////////////////////////////////////////////////////gYHg/ysrzP+GhuL/+vr9///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////Q0PP/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/zg4z//6+v3//////5mZ5v8sLMz/KyvM/ysrzP9HR9L/7u77////////////////////////////////////////////////////////////ra3r/ysrzP8rK8z/KyvM/1dX1v/u7vv/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////aGja/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/y8vy//////+np+n/KyvM/ysrzP8rK8z/KyvM/2Fh2f/7+/7/////////////////////////////////////////////////xcXx/y4uzP8rK8z/KyvM/ysrzP9ZWdf//f3+/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////f3+/zMzzv8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/83N8////////f3+/1xc2P8rK8z/KyvM/ysrzP8rK8z/h4fi////////////////////////////////////////////1NT0/zU1zv8rK8z/KyvM/ysrzP8xMc3/3t73/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+bm+f8rK8z/KyvM/ysrzP8rK8z/fX3f/6mp6v9ZWdf/KyvM/0pK0//7+/7////////////j4/j/ODjP/ysrzP8rK8z/KyvM/ysrzP/Dw/D/////////////////////////////////5ub5/z090P8rK8z/KyvM/ysrzP8rK8z/n5/o///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////s7Pr/KyvM/ysrzP8rK8z/NTXO//v7/v////////////Dw+//4+P3//////////////////////8DA8P8sLMz/KyvM/ysrzP8rK8z/RUXS//r6/f///////////////////////////15e2P8rK8z/KyvM/ysrzP8rK8z/dHTd/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////z8/0f8rK8z/KyvM/y4uzP/6+v3/////////////////////////////////////////////////o6Pp/ysrzP8rK8z/KyvM/ysrzP+zs+3//////////////////////9jY9f8rK8z/KyvM/ysrzP8rK8z/V1fW//j4/f////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////9mZtr/KyvM/ysrzP8rK8z/1dX1//////////////////////////////////////////////////////+np+n/LCzM/ysrzP8rK8z/t7ft///////////////////////4+P3/OzvQ/ysrzP8rK8z/YGDY/+zs+v//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////jIzj/ysrzP8rK8z/KyvM/7Gx7P///////////////////////////////////////////////////////////8TE8f9padv/oqLp//39/v///////////////////////////9zc9v+Tk+X/uLju//v7/v///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////7Oz7f8rK8z/KyvM/ysrzP+JieL////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////Y2PX/KyvM/ysrzP8rK8z/Y2PZ////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+/v+/zg4z/8rK8z/KyvM/zExzf/4+P3////////////////////////////////////////////////////////////f3/f/oqLo/3Jy3f9gYNj/SUnT/0VF0v9FRdL/RUXS/09P1f9kZNn/k5Pl/+Hh9/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////9wcNz/KyvM/ysrzP8rK8z/ysry/////////////////////////////////////////////////+Pj+P9qatv/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8uLsz/np7n////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////rKzr/ysrzP8rK8z/KyvM/5GR5P///////////////////////////////////////////8rK8v84OM//KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP+MjOP//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+fn+f8rK8z/KyvM/ysrzP9ZWdf//////////////////////////////////////+fn+f87O9D/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/QEDR/0VF0v9FRdL/RUXS/z090P8rK8z/KyvM/ysrzP8rK8z/Li7M/9nZ9v//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////UVHV/ysrzP8rK8z/Li7M//Hx+/////////////////////////////////9/f+D/KyvM/ysrzP8rK8z/KyvM/2ho2//BwfD/8fH7////////////////////////////2dn2/0VF0v8rK8z/KyvM/ysrzP9yct3//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////4yM4/8rK8z/KyvM/ysrzP+7u+7////////////////////////////x8fv/MzPO/ysrzP8rK8z/KyvM/5CQ5P/////////////////////////////////////////////////Q0PP/KyvM/ysrzP8rK8z/ODjP//v7/v////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+5ue7/KyvM/ysrzP8rK8z/fX3f////////////////////////////w8Pw/ysrzP8rK8z/KyvM/0lJ0//9/f7//////////////////////////////////////////////////////0xM1P8rK8z/KyvM/ysrzP/Pz/P/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////ubnu/9zc9v//////////////////////wMDv/ysrzP8rK8z/KyvM/1tb1////////////////////////////5WV5f8rK8z/KyvM/ysrzP+VleX///////////////////////////////////////////////////////////9/f+D/KyvM/ysrzP8rK8z/sbHs/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////2tr2/8sLMz/a2vb/8bG8f/9/f7/39/3/1tb1/8rK8z/KyvM/ysrzP9YWNf///////////////////////////99fd//KyvM/ysrzP8rK8z/xcXx////////////////////////////////////////////////////////////m5vn/ysrzP8rK8z/KyvM/5+f6P////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////9UVNb/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/fHzf////////////////////////////cHDc/ysrzP8rK8z/KyvM/+Li+P///////////////////////////////////////////////////////////7S07f8rK8z/KyvM/ysrzP+OjuT/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////PT3Q/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/76+7////////////////////////////2Nj2f8rK8z/KyvM/ysrzP/s7Pr///////////////////////////////////////////////////////////+9ve//KyvM/ysrzP8rK8z/h4fi////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+vr9/zAwzf8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/29v3P////////////////////////////////9ra9v/KyvM/ysrzP8rK8z/5ub4////////////////////////////////////////////////////////////vb3v/ysrzP8rK8z/KyvM/4eH4v/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////m5vj/hobh/z090P8rK8z/KyvM/ysrzP8rK8z/KyvM/3V13v/19fz/////////////////////////////////eHje/ysrzP8rK8z/KyvM/87O8////////////////////////////////////////////////////////////7297/8rK8z/KyvM/ysrzP+MjOP////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////7+/7/zc3z/5ub5/+VleX/n5/o/93d9////////////////////////////////////////////4SE4f8rK8z/KyvM/ysrzP+2tu3///////////////////////////////////////////////////////////+9ve//KyvM/ysrzP8rK8z/lZXl//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////39/v/U1PT/m5vn/6Ki6P/k5Pj///////////////////////////////////////////+Tk+X/KyvM/ysrzP8rK8z/n5/o////////////////////////////////////////////////////////////uLju/ysrzP8rK8z/KyvM/5qa5v///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////76+7/9SUtX/KyvM/ysrzP8rK8z/KyvM/3x83//w8Pv/////////////////////////////////trbt/ysrzP8rK8z/KyvM/3V13f///////////////////////////////////////////////////////////52d5/8rK8z/KyvM/ysrzP+vr+z/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////+vr9/3x83/8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/OzvQ/9DQ8////////////////////////////9jY9f8rK8z/KyvM/ysrzP87O9D//f3+//////////////////////////////////////////////////////9zc93/KyvM/ysrzP8rK8z/ysry/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////21t3P8rK8z/KyvM/ysrzP9mZtr/pKTp/5WV5f9ISNP/KyvM/ysrzP81Nc7/xsbx///////////////////////6+v3/MzPO/ysrzP8rK8z/KyvM/9TU9P/////////////////////////////////////////////////19fz/ODjP/ysrzP8rK8z/KyvM/+vr+v///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////6Oj6f8rK8z/KyvM/0RE0v/ExPH/////////////////+vr9/5aW5v8sLMz/KyvM/zs70P/u7vv//////////////////////2Zm2v8rK8z/KyvM/ysrzP+EhOH/8fH7////////////////////////////////////////////lZXl/ysrzP8rK8z/KyvM/0xM1P////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////39/v8/P9H/KyvM/z090P/m5vn/////////////////////////////////p6fp/ysrzP8rK8z/enrf//////////////////////+goOj/KyvM/ysrzP8rK8z/KyvM/ysrzP9SUtX/gYHg/5WV5f+np+r/r6/s/6+v7P+Tk+X/VlbW/ysrzP8rK8z/KyvM/ysrzP+QkOT////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////Y2PX/KyvM/ysrzP+dnef////////////4+P3/mJjm/5iY5v/6+v3///////////9hYdn/KyvM/ywszP/i4vj/////////////////7u77/zAwzf8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8wMM3/5+f5////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////vr7v/ysrzP8rK8z/y8vy////////////lpbm/ysrzP8rK8z/rq7r////////////0NDz/ysrzP8rK8z/kJDk//////////////////////99fd//KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/ysrzP8uLsz/ubnu/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////8bG8f8rK8z/KyvM/6en6v///////f3+/1RU1v8rK8z/KyvM/3d33v////////////////87O9D/KyvM/1RU1v//////////////////////6+v6/z8/0f8rK8z/LCzM/3d33v9UVNb/LCzM/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/zU1zv9vb9z/2tr2///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////z8/z/NTXO/ysrzP84OM//nZ3n/3Bw3P8rK8z/KyvM/ysrzP9CQtH/9fX8///////9/f7/OjrP/ysrzP8/P9D////////////////////////////f3/f/fX3g/7m57v////////////j4/f/h4ff/ysry/8rK8v/KyvL/ysry/+Li+P/7+/7//////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////52d5/8rK8z/KyvM/ysrzP8rK8z/KyvM/2tr2/8rK8z/KyvM/3V13v/f3/f/mJjm/ysrzP8rK8z/VlbW///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////6+v3/cHDd/ysrzP8rK8z/KyvM/15e2P/19fz/V1fW/ysrzP8rK8z/KyvM/ysrzP8rK8z/KyvM/7a27f/////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////NzfP/r6/s/8XF8f/9/f7//////+Pj+P9HR9L/KyvM/ysrzP8rK8z/KyvM/3p63///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////6+v6/5OT5f9jY9n/cnLd/8DA7///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////////AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA='
gif64en = r'R0lGODlhgACAAPUgAMwrK84yMs84ONE/P9NGRtRMTNZTU9daWtlgYNtnZ9xubt50dOB7e+GCguOIiOSPj+WVleebm+iiouqoqOyvr+22tu+8vPDDw/LKyvTQ0PXX1/fe3vnk5Prr6/zy8v34+P///wAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACH5BAAAAAAAIf8LSW1hZ2VNYWdpY2sNZ2FtbWE9MC40NTQ1NQAsAAAAAIAAgAAABv5AkHBILBqPyKRyyWw6n9CodEqtWq/YrHbL7Xq/4LB4TC6bz+i0es1uu9/wuHxOr9vv+Lx+zy9uJBATGH18FAQAiIgDF4R4EYmQiAwejWkYCggICRQfRA+RoAaVZh4JoAMaQhagrBKjYx0DrAACqQWzkQOdr2AMuAAGGpEFFRgQoBC8Xxm/iAeQCbsgEpEKyl0fBpEBCAKzAalDHwGQBUwcFtdOwpAIHCAdC6wNRgiQB0ofDogCFEUcGCxAGEhw4AUMlPpQSxSgg7hbkfwVeZYogRJfDC0YS0CuGSIDD8LlwYiIHhEKoEQOkZUoGRKUHmOWdIiHJACXQzhsM8Iukf5KIuNkCq2FpwGkCEYgInJgRAGkAUkWCpVJ1I5UANb+HQJGc0hPRK6QsJwqU5edC5E2HPHwzghFRA2RrCI7dYKdoIkWPJkQiSkSe0IHZBo8dpYoO/KONrnQUZG0Ih4as0IAAUPbIxwoaJvFqA6zSA8eG6EgGUA6ubMqR7kKyW+dY8MyHOlgtJqSfaGo4I502A7gaiI3QPAWiUDXI0pbU/kKKWGdWN8QJIdk4LiRDrMGUfkN6bSdDm9lHnCOBC2oAKKhsAab50PimArSG1kPwGIV5kv3VNiKa5GT2pGEVQUrWe1hyCyhPRFeImpZMR0whHxwwQObKTCBdUwQB0kAWP6YkssrHpD3BHagmHOFTYmoE4VOoCCABYqI8OFBBQ3wBwwElzWBHwAMvMhKexW8NwsDGCKBASs4VQEjAPK98YEEhf1CwE/lsWLXiazkGAd0UwlQZBEwRdKZFUtqGQddWDHBV0o+gmImHDYK5R0SsEXyZhRl2rFgWUvQJxuWbiKGJgANvsSKdlYImcidbuxG1pxGzCUmFtylaFUkCVTQyR9RJvKAEkeCImAVcdJyRwVwMVDoEBmUJokSO+p1BSsE3NGqA1+C4BQoJiFB4jBX/NrOHR/kKkSdkCR5RKeEWiEpNMo4mkgFS+wKmhXSlqRMpYgYOw0rZlGxWSSQNqJhIv74LBEqKD1OsSMAIhLyWSTKIvEgIkhJYe2wvCCbCKJK0IdIAowe4W9LyuwpgBMfnAvKARQAjBmAoKzaiAcEPiFwiYN1jACzibT7SpiQjMrEvoPO4qUyS1KphAf3ppyIyZU4DECtUXTArcz1XbOjyFF8wjPC1xwMgETuLtlluaPsXHATHUCwczMFROCtjLRykQEgEGDicSYOBOIyL+b1pSIc9CF9Nhstr+3GuJA06XYZrvZWxQcCPQ1UvMrEekUGSjEgtxAcPOANSHxXgioo+VLBWIDVDjM2IUKTWwXJiUCVBIvbUKuMh5FcbWjWUX0jcSOdLjzF4qy4dkTZoAigt/4dGOcmRavfeAv3MInnAXtzRXxgWRIdlOopE/OyYl8jzyY7hAXj0izE1DcPPsQjuLjOxwa4CIBAnAHEWzkrfzahaCSnt0eX2iAkjwwUH4AOCgHW0xFz9kDd72IUHuwOSa98OF8zAGg0RBCgd0ooHi7Sd4d34QJpGnAVPywGBQ7YzID1kwPKmtGVAprmbxL8YB8c2KIiTMBVjbMC6zLHQDwAziMJiBcGziWrLIwPAaKrwweUBowU/mMzBsggwzbzKXVooAIFocDVPKCAAcwOChrwz9yk8MQoCHGKWMyiFrfIxS568YtgDKMYx0jGMprxjGhMoxrXyMY2uvGNcIyjHBnnSMc62vGOeMyjHvfIxz768Y+ADKQguRgEADs='

#V1.0 Transfer Module, Stopped!
def socket_client_send_data(data):
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.setsockopt(socket.SOL_SOCKET, socket.SO_KEEPALIVE, 1)
        s.connect(('www.lijishi.cn', 6666))
        #s.listen(10)
        #s.connect(('127.0.0.1', 6666))  #Local test
    except socket.error as msg:
        if str(msg) == '[WinError 10061] 由于目标计算机积极拒绝，无法连接。':
            InternetError()
        print(msg)
        print(sys.exit(1))
    s.send(data.encode())  #将要传输的数据编码发送，如果是字符数据就必须要编码发送

    #sock = s.accept()
    buf = s.recv(1024)  # 接收数据
    buf = buf.decode()
    s.close()
    if str(buf)[0:1] == '1':
        #print('Yes!')
        #print(str(buf)[2:])
        return str(buf)
    elif str(buf)[0:1] == '2':
        #print('Left!')
        #print(str(buf)[2:])
        return str(buf)[2:]
    else:
        #print('Error!')
        #print(str(buf)[2:])
        return str(buf)

#V2.0 Transfer Module, Stopped!
def api_client_send_data(data):
    param = data.split('+')
    baseapiurl = 'http://api.sometools.online/release/ocr/'
    if param[0] == '00':
        mode = param[1]
        apiurl = baseapiurl + '1+' + str(mode) + '/'
        response = requests.get(url=apiurl)
        result = response.content.decode("utf-8")
        if int(result) < 1000:
            return 1000-int(result)
        else:
            return 0
    else:
        mode = param[0]
        key = param[1]
        apiurl = baseapiurl + '2+' + str(data) + '/'
        response = requests.get(url=apiurl)
        result = response.content.decode("utf-8")
        return result


def Icon():
    '''
    with open("mei.ico", "rb") as icon:
        icon64en = base64.b64encode(icon.read())
        print(icon64en)
    '''
    global  icon64en
    with open('tmp.ico', 'wb') as tmp:
        tmp.write(base64.b64decode(icon64en))

def Gif():
    '''
    with open("mei.gif", "rb") as gif:
        gif64en = base64.b64encode(gif.read())
        print(gif64en)
    '''
    global  gif64en
    with open('temp.gif', 'wb') as temp:
        temp.write(base64.b64decode(gif64en))
    mei_gif = tk.PhotoImage(file="temp.gif")

def Base64EnCode(path):
    with open(path, "rb") as i:
        input64en = base64.b64encode(i.read())
    return input64en

def Base64DeCode_JPG(str_in):
    global  path_out
    string = str(path_in.get())
    string = string.replace("\\", "\\\\")
    string = string.replace("/", "\\\\")
    i = string.rfind("\\\\")
    if string[-3:].lower() == 'jpg' or string[-3:].lower() == 'png':
        out_name_temp = string[i + 2:len(string) - 4]
    elif string[-4:].lower() == 'jpeg':
        out_name_temp = string[i + 2:len(string) - 5]
    path = str(path_out.get()) + out_name_temp + '_OCR_' + str(int(time.time())) + '.jpg'
    with open(path, 'wb+') as out:
        out.write(base64.b64decode(str_in))
    img = Image.open(path)
    out_pic = ImageTk.PhotoImage(img)

def Base64DeCode_XLSX(xlsx_str):
    string = str(path_in.get())
    string = string.replace("\\", "\\\\")
    string = string.replace("/", "\\\\")
    i = string.rfind("\\\\")
    if string[-3:].lower() == 'jpg' or string[-3:].lower() == 'png':
        out_name_temp = string[i + 2:len(string) - 4]
    elif string[-4:].lower() == 'jpeg':
        out_name_temp = string[i + 2:len(string) - 5]
    out_name = str(path_out.get()) + out_name_temp + "_OCR_" + str(int(time.time())) + ".xlsx"
    with open(out_name, 'wb') as out:
        out.write(base64.b64decode(xlsx_str))

def SelectPath_In():
    path_ = tkinter.filedialog.askopenfilename()
    path_ = path_.replace("/", "\\\\")
    path_in.set(path_)
    string = path_
    i = string.rfind("\\\\")
    string = string[:i+2]
    path_out.set(string)

def SelectPath_Out():
    path_ = tkinter.filedialog.askdirectory()
    path_ = path_.replace("/", "\\\\")
    path_out.set(path_)

def ModeQuery(mode):
    '''
    11:通用印刷体识别 12:通用印刷体识别（高速版）13:通用印刷体识别（高精度版） 14:通用手写体识别 15:英文识别
    21:身份证识别 22:营业执照识别 23:银行卡识别 24:名片识别
    31: 增值税发票识别 32：运单识别
    41：行驶证/驾驶证识别 42：车牌识别 43：车辆VIN码识别
    51：算式识别 52：表格识别
    '''
    if mode == '11':
        mode_data = '  --通用印刷体识别'
    elif mode == '12':
        mode_data = '  --通用印刷体识别（高速版）'
    elif mode == '13':
        mode_data = '  --通用印刷体识别（高精度版）'
    elif mode == '14':
        mode_data = '  --通用手写体识别'
    elif mode == '15':
        mode_data = '  --英文识别'
    elif mode == '21':
        mode_data = '  --身份证识别'
    elif mode == '22':
        mode_data = '  --营业执照识别'
    elif mode == '23':
        mode_data = '  --银行卡识别'
    elif mode == '24':
        mode_data = '  --名片识别'
    elif mode == '31':
        mode_data = '  --增值税发票识别'
    elif mode == '32':
        mode_data = '  --运单识别'
    elif mode == '41':
        mode_data = '  --驾驶证识别'
    elif mode == '42':
        mode_data = '  --车牌识别'
    elif mode == '43':
        mode_data = '  --车辆VIN码识别'
    elif mode == '44':
        mode_data = '  --行驶证识别'
    elif mode == '51':
        mode_data = '  --算式识别'
    elif mode == '52':
        mode_data = '  --表格识别'
    else:
        mode_data = '  --未知类型'
    return mode_data

def treeviewClick(event):
    item = mode_tree.selection()[0]
    mode = str(mode_tree.item(item, "values"))[2:4]
    mode_temp = mode
    if mode[1:] == '\'':
        return
    else:
        if mode_temp == '44':
            mode_temp = '41'
        data = '00+' + mode_temp
        #return_data = socket_client_send_data(data)
        return_data = api_client_send_data(data)
    mode_data = ModeQuery(mode)
    left_data = str(return_data) + mode_data
    left.set(left_data)

def Tip():
    tk.messagebox.showinfo("Tips", "图像格式仅支持jpg/jpeg/png\n图像大小上限未知，推荐1MB以下\n输出格式不选择则为不输出文本\n输出的.doc文件请以utf-8编码打开\n表格识别默认在输出位置生成.xlsx文件\n识别结果来自于腾讯云-文字识别API\n请勿用于非法用途，作者不负任何法律责任\n腾讯云若更改免费政策，可能导致软件无法使用")

def UsingKeyTip():
    tk.messagebox.showinfo("Tips", "Q：何为UsingKey？\nA：一组多位字母+数字混合密钥\nQ：如何获取UsingKey？\nA：目前仅可通过联系作者获取\nQ：没有UsingKey还想用？\nA：请根据教程自行通过API使用")

def Finish():
    tk.messagebox.showinfo("Finish", "识别完成！")

def InternetError():
    tk.messagebox.showerror("Error", "网络连接失败！\n请确认所在网络可连接性后重试\n若确认后依然报错，请联系作者")

def UsingKeyError():
    tk.messagebox.showerror("Error", "UsingKey校验失败！\n请确认UsingKey正确性后重试\n若确认后依然报错，请联系作者")

def IdentifyError():
    tk.messagebox.showerror("Error", "文本识别失败！\n所选服务本月可使用次数已尽\n请更换服务类型或下月重试")

def ParamError():
    tk.messagebox.showerror("Error", "缺少必要参数\n请检查后重试")

def UnknownError():
    tk.messagebox.showerror("Error", "未知错误\n请确认各项后重试\n若确认后依然报错，请联系作者")

def Start():
    '''
    try:
        if path_in.get()[-3:] == "jpg":
            pic = 1
        elif path_in.get()[-4:] == "jpeg":
            pic = 1
        elif path_in.get()[-3:] == "png":
            pic = 0
        pic_str = Base64EnCode(str(path_in.get()))
    except:
        Error()
    '''
    global option
    global idcard_cut
    global picture_cut
    global out_format

    pic_str = Base64EnCode(str(path_in.get()))
    usingkey = using_key.get()
    item = mode_tree.selection()[0]
    mode = str(mode_tree.item(item, "values"))[2:4]
    mode_temp = mode

    if mode[1:] == '\'':
        UnknownError()
    else:
        if mode_temp == '44':
            mode_temp = '41'
        data = mode_temp + '+' + str(usingkey)
        #return_data = socket_client_send_data(data)
        return_data = api_client_send_data(data)

    if return_data[1:2] == '1':
        mode_data = ModeQuery(mode)
        left_data = return_data[3:-1] + mode_data
        left.set(left_data)

        business_card_flag = 0
        if mode == '21':
            IsFront = int(option.get()) - 1 #0正1背
            IsCardCut = int(idcard_cut.get()) - 1 #0是1否
            IsPicCut = int(picture_cut.get()) - 1 #0是1否
            return_text = Tencent_OCR_IDCard(str(pic_str)[2:-1], IsFront, IsCardCut, IsPicCut)

        elif mode == '24':
            if business_card.get() == "是":
                business_card_flag = 1
            else:
                business_card_flag = 0
            return_text = Tencent_OCR_BusinessCard(str(pic_str)[2:-1], business_card_flag)
            out_text = Text_Arrange_Summary(return_text, mode)

        elif mode == '44':
            if vehicle_license.get() == "行驶证主页正面":
                vehicle_license_flag = 0
            elif vehicle_license.get() == "行驶证副页正面":
                vehicle_license_flag = 1
            else:
                ParamError()
            return_text = Tencent_OCR_VehicleLicense(str(pic_str)[2:-1], vehicle_license_flag)
            vehicle_license_warning = Text_Arrange_VehicleLicense_Warning(return_text)
        else:
            return_text = Tencent_OCR_Summary(str(pic_str)[2:-1], mode)

        out_text = Text_Arrange_Summary(return_text, mode)

        if business_card_flag:
            out_pic = Text_Arrange_BusinessCard_Picture(return_text)
            Base64DeCode_JPG(out_pic)

        if mode == '21':
            idcard_warning = Text_Arrange_IDCard_Warning(return_text)
            if IsFront:
                if IsCardCut == 0:
                    out_pic = Text_Arrange_IDCard_IDCardCut(return_text)
                    Base64DeCode_JPG(out_pic)
            else:
                if IsCardCut == 0:
                    out_pic = Text_Arrange_IDCard_IDCardCut(return_text)
                    Base64DeCode_JPG(out_pic)
                    time.sleep(1)
                if IsPicCut == 0:
                    out_por = Text_Arrange_IDCard_PortraitCut(return_text)
                    Base64DeCode_JPG(out_por)

        if mode == '41':
            driver_license_warning = Text_Arrange_DriverLicense_Warning(return_text)

        if mode == '52':
            Base64DeCode_XLSX(out_text)
            scr.delete(1.0, END)
            scr.insert("insert", "表格识别模式无文字预览\n")
            scr.insert("insert", "请移步至输出位置查看Excel文件\n")
            Finish()
            return

        scr.delete(1.0, END)
        if business_card_flag:
            scr.insert("insert", "修正后图片已在输出位置生成\n")
            scr.insert("insert", "\n")

        if mode == '21':
            if IsCardCut == 0:
                scr.insert("insert", "裁剪后图片已在输出位置生成\n")
            if IsPicCut== 0:
                scr.insert("insert", "头像图片已在输出位置生成\n")
            if IsCardCut == 0 or IsPicCut== 0:
                scr.insert("insert", "\n")
            if idcard_warning != 'null':
                scr.insert("insert", "！！！Warning！！！告警信息！！！\n")
                scr.insert("insert", idcard_warning)
                scr.insert("insert", "\n")

        if mode == '41':
            if driver_license_warning != 'null':
                scr.insert("insert", "！！！Warning！！！告警信息！！！\n")
                scr.insert("insert", driver_license_warning)
                scr.insert("insert", "\n")
        if mode == '44':
            if vehicle_license_warning != 'null':
                scr.insert("insert", "！！！Warning！！！告警信息！！！\n")
                scr.insert("insert", vehicle_license_warning)
                scr.insert("insert", "\n")
                if vehicle_license_flag:
                    scr.insert("insert", "行驶证副页正面\n")
                    scr.insert("insert", "\n")
                else:
                    scr.insert("insert", "行驶证主页正面\n")
                    scr.insert("insert", "\n")
        scr.insert("insert", out_text)

        string = str(path_in.get())
        string = string.replace("\\", "\\\\")
        string = string.replace("/", "\\\\")
        i = string.rfind("\\\\")
        if string[-3:].lower() == 'jpg' or string[-3:].lower() == 'png':
            out_name_temp = string[i+2:len(string)-4]
        elif string[-4:].lower() == 'jpeg':
            out_name_temp = string[i+2:len(string)-5]

        if int(out_format.get()) == 1:
            out_name = str(path_out.get()) + out_name_temp + "_OCR_" + str(int(time.time())) + ".txt"
        elif int(out_format.get()) == 2:
            out_name = str(path_out.get()) + out_name_temp + "_OCR_" + str(int(time.time())) + ".doc"
        else:
            Finish()
            return

        if int(out_format.get()) == 1: #TXT
            with open(out_name, 'w', encoding='utf-8') as out:
                now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
                #out.write("##############################\n")
                out.write("----------------------------------------\n")
                out.write("Created By OCR Lite At ")
                out.write(now_time)
                out.write("\n")
                out.write("OCR Mode : ")
                out.write(ModeQuery(mode)[4:])
                out.write("\n")
                out.write("----------------------------------------\n")
                #out.write("##############################\n")
                out.write("\n")

                if mode == '44':
                    if vehicle_license_warning != 'null':
                         out.write("！！！Warning！！！告警信息！！！\n")
                         out.write(vehicle_license_warning)
                         out.write("\n")
                    if vehicle_license_flag:
                         out.write("行驶证副页正面\n")
                         out.write("\n")
                    else:
                         out.write("行驶证主页正面\n")
                         out.write("\n")

                if mode == '21':
                    if driver_license_warning != 'null':
                        out.write("！！！Warning！！！告警信息！！！\n")
                        out.write(idcard_warning)
                        out.write("\n")

                if mode == '41':
                    if driver_license_warning != 'null':
                        out.write("！！！Warning！！！告警信息！！！\n")
                        out.write(driver_license_warning)
                        out.write("\n")

                out.write(out_text)
                out.write("\n")
            out.close()

        else: #DOC
            now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
            doc = Document()
            doc.styles['Normal'].font.name = u'宋体'
            doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("---------------------------------------------\n")
            run = paragraph.add_run("Created By OCR Lite At ")
            run = paragraph.add_run(now_time)
            run = paragraph.add_run("\n")
            run = paragraph.add_run("OCR Mode : ")
            run = paragraph.add_run(ModeQuery(mode)[4:])
            run = paragraph.add_run("\n")
            run = paragraph.add_run("---------------------------------------------\n")
            run = paragraph.add_run("\n")

            if mode == '44':
                if vehicle_license_warning != 'null':
                    run = paragraph.add_run("！！！Warning！！！告警信息！！！\n")
                    run = paragraph.add_run(vehicle_license_warning)
                    run = paragraph.add_run("\n")
                if vehicle_license_flag:
                    run = paragraph.add_run("行驶证副页正面\n")
                    run = paragraph.add_run("\n")
                else:
                    run = paragraph.add_run("行驶证主页正面\n")
                    run = paragraph.add_run("\n")

            if mode == '21':
                if idcard_warning != 'null':
                    run = paragraph.add_run("！！！Warning！！！告警信息！！！\n")
                    run = paragraph.add_run(idcard_warning)
                    run = paragraph.add_run("\n")

            if mode == '41':
                if driver_license_warning != 'null':
                    run = paragraph.add_run("！！！Warning！！！告警信息！！！\n")
                    run = paragraph.add_run(driver_license_warning)
                    run = paragraph.add_run("\n")

            run = paragraph.add_run(out_text)
            run = paragraph.add_run("\n")
            font = run.font
            doc.save(out_name)

        Finish()
    else:
        if return_data[3:-1] == '101':
            UsingKeyError()
        elif return_data[3:-1] == '102':
            IdentifyError()

def TK_IDCard():
    global option
    global idcard_cut
    global picture_cut

    # window centered
    idcard_window = Toplevel()
    screen_width = idcard_window.winfo_screenwidth()
    screen_heigh = idcard_window.winfo_screenheight()
    idcard_window_width = 215
    idcard_window_heigh = 240
    x = (screen_width - idcard_window_width) / 2
    y = (screen_heigh - idcard_window_heigh) / 2
    idcard_window.geometry("%dx%d+%d+%d" % (idcard_window_width, idcard_window_heigh, x, y))

    # window layout
    idcard_window.title('Params')
    Icon()
    idcard_window.iconbitmap('tmp.ico')
    os.remove('tmp.ico')

    option = IntVar()
    option_frame = ttk.LabelFrame(idcard_window, text='身份证双面选择')
    option_frame.grid(row=0, column=0, rowspan=1, columnspan=4, padx=10, pady=5)
    ttk.Radiobutton(option_frame, text="人像面", variable=option, value=1).grid(row=1, column=1, padx=10, pady=5)
    ttk.Radiobutton(option_frame, text="国徽面", variable=option, value=2).grid(row=1, column=2, padx=10, pady=5)

    picture_cut = IntVar()
    picture_frame = ttk.LabelFrame(idcard_window, text='人像照片裁剪')
    picture_frame.grid(row=2, column=0, rowspan=1, columnspan=4, padx=10, pady=5)
    ttk.Radiobutton(picture_frame, text="是", variable=picture_cut, value=1).grid(row=3, column=1, padx=10, pady=5)
    ttk.Radiobutton(picture_frame, text="否", variable=picture_cut, value=2).grid(row=3, column=2, padx=10, pady=5)

    idcard_cut = IntVar()
    idcard_frame = ttk.LabelFrame(idcard_window, text='身份证照片裁剪')
    idcard_frame.grid(row=4, column=0, rowspan=1, columnspan=4, padx=10, pady=5)
    ttk.Radiobutton(idcard_frame, text="是", variable=idcard_cut, value=1).grid(row=5, column=1, padx=10, pady=5)
    ttk.Radiobutton(idcard_frame, text="否", variable=idcard_cut, value=2).grid(row=5, column=2, padx=10, pady=5)

    ttk.Button(idcard_window, width=25, text="开始", command=Start).grid(row=6, column=0, padx=10, pady=5)

def About():
    # window centered
    about_window = Toplevel()
    screen_width = about_window.winfo_screenwidth()
    screen_heigh = about_window.winfo_screenheight()
    about_window_width = 400
    about_window_heigh = 240
    x = (screen_width - about_window_width) / 2
    y = (screen_heigh - about_window_heigh) / 2
    about_window.geometry("%dx%d+%d+%d" % (about_window_width, about_window_heigh, x, y))

    # window layout
    global shi_gif
    about_window.title('About')
    Icon()
    Gif()
    about_window.iconbitmap('tmp.ico')
    os.remove('tmp.ico')
    shi_gif = tk.PhotoImage(file="temp.gif")
    #os.remove('temp.gif')
    software_frame = ttk.LabelFrame(about_window, text='Software Info')
    software_frame.grid(row=0, column=0, rowspan=5, columnspan=4, padx=50, pady=5)
    ttk.Label(software_frame, image=shi_gif, compound='left').grid(row=0, rowspan=3, column=0)
    os.remove('temp.gif')
    ttk.Label(software_frame, text="OCR Lite 2.0").grid(row=0, column=1, sticky = W)
    ttk.Label(software_frame, text="@Author    :   lijishi").grid(row=1, column=1, sticky = W)
    ttk.Label(software_frame, text="@EditTime  :   Oct 6,2020").grid(row=2, column=1, sticky = W)

    copyright_frame = ttk.LabelFrame(about_window, text='LICENSE Info')
    copyright_frame.grid(row=5, column=0, rowspan=3, columnspan=4, padx=50, pady=5)
    ttk.Label(copyright_frame, text = "Github @ OCR_TencentCloudAPI").grid(row=5, column=0)
    ttk.Label(copyright_frame, text="GNU GENERAL PUBLIC LICENSE Version 3").grid(row=6, column=0)

# window centered
main_window=tk.Tk()
screen_width = main_window.winfo_screenwidth()
screen_heigh = main_window.winfo_screenheight()
main_window_width = 520
main_window_heigh = 435
x = (screen_width-main_window_width) / 2
y = (screen_heigh-main_window_heigh) / 2
main_window.geometry("%dx%d+%d+%d" %(main_window_width,main_window_heigh,x,y))

# window layout
global mei_gif
main_window.title("OCR Lite V2.0")
Icon()
main_window.iconbitmap('tmp.ico')
os.remove('tmp.ico')

path_frame = ttk.LabelFrame(main_window, text = '路径选择')
path_frame.grid(row = 0, column = 0, rowspan = 2, columnspan = 6, padx=5, pady=5)
path_in = tk.StringVar()
path_out = tk.StringVar()
path_in.set("请选择待识别图片位置，也可在框中键入")
path_out.set("请选择输出文档位置，默认为源文件路径")
ttk.Label(path_frame, text = "待识别图片位置").grid(row = 0, column = 0, padx=5)
ttk.Entry(path_frame, width = 40, textvariable = path_in).grid(row = 0, column = 1, padx=5)
ttk.Button(path_frame, text = "选择", command = SelectPath_In).grid(row = 0, column = 2, padx=5)
ttk.Label(path_frame, text = "输出文件位置").grid(row = 1, column = 0, padx=5)
ttk.Entry(path_frame, width = 40, textvariable = path_out).grid(row = 1, column = 1, padx=5)
ttk.Button(path_frame, text = "选择", command = SelectPath_Out).grid(row = 1, column = 2, padx = 5, pady = 5)

option_frame = ttk.LabelFrame(main_window, text = '模式选择')
option_frame.grid(row = 2, column = 0, padx=5, pady=5)
mode_tree = ttk.Treeview(option_frame, height=5, show = "tree")
mode_tree_F1 = mode_tree.insert("", 0, text="通用文字识别", values=("1"))
mode_tree_F2 = mode_tree.insert("", 1, text="卡证文字识别", values=("2"))
mode_tree_F3 = mode_tree.insert("", 2, text="票据单据识别", values=("3"))
mode_tree_F4 = mode_tree.insert("", 3, text="汽车相关识别", values=("4"))
mode_tree_F5 = mode_tree.insert("", 4, text="行业文档识别", values=("5"))
mode_tree_F1_1 = mode_tree.insert(mode_tree_F1, 0, text="通用印刷体识别", values=("11"))
mode_tree_F1_2 = mode_tree.insert(mode_tree_F1, 1, text="通用印刷体识别（高速版）", values=("12"))
mode_tree_F1_3 = mode_tree.insert(mode_tree_F1, 2, text="通用印刷体识别（高精度版）", values=("13"))
mode_tree_F1_4 = mode_tree.insert(mode_tree_F1, 3, text="通用手写体识别", values=("14"))
mode_tree_F1_5 = mode_tree.insert(mode_tree_F1, 4, text="英文识别", values=("15"))
mode_tree_F2_1 = mode_tree.insert(mode_tree_F2, 0, text="身份证识别", values=("21"))
mode_tree_F2_2 = mode_tree.insert(mode_tree_F2, 1, text="营业执照识别", values=("22"))
mode_tree_F2_3 = mode_tree.insert(mode_tree_F2, 2, text="银行卡识别", values=("23"))
mode_tree_F2_4 = mode_tree.insert(mode_tree_F2, 3, text="名片识别", values=("24"))
mode_tree_F3_1 = mode_tree.insert(mode_tree_F3, 0, text="增值税发票识别", values=("31"))
mode_tree_F3_2 = mode_tree.insert(mode_tree_F3, 1, text="运单识别", values=("32"))
mode_tree_F4_1 = mode_tree.insert(mode_tree_F4, 0, text="驾驶证识别", values=("41"))
mode_tree_F4_4 = mode_tree.insert(mode_tree_F4, 1, text="行驶证识别", values=("44"))
mode_tree_F4_2 = mode_tree.insert(mode_tree_F4, 2, text="车牌识别", values=("42"))
mode_tree_F4_3 = mode_tree.insert(mode_tree_F4, 3, text="车辆VIN码识别", values=("43"))
mode_tree_F5_1 = mode_tree.insert(mode_tree_F5, 0, text="算式识别", values=("51"))
mode_tree_F5_2 = mode_tree.insert(mode_tree_F5, 1, text="表格识别", values=("52"))
#item = mode_tree.selection()[0]
#mode_tree.item(item, "values")[2:4]
yscrollbar = Scrollbar(option_frame, orient=VERTICAL, command=mode_tree.yview)
mode_tree.configure(yscrollcommand = yscrollbar.set)
yscrollbar.pack(side = RIGHT, fill = Y)
mode_tree.bind('<ButtonRelease-1>', treeviewClick)
mode_tree.pack(side=LEFT)

text_frame = ttk.LabelFrame(main_window, text = '输出文本')
text_frame.grid(row = 1, column = 1, rowspan = 2, columnspan = 2, pady=5, sticky = S)
#ttk.Entry(text_frame, width = 28, textvariable = text_out).grid(row = 3, column = 1, rowspan = 2, padx=5, pady=5)
#text = tk.Text(text_frame, width = 29, height = 7).grid(row = 3, column = 1, columnspan = 2, padx=5, pady=5)
#scr.insert("insert", "python")
#scr.delete(1.0, END)
scr = scrolledtext.ScrolledText(text_frame, width = 33, height = 7, wrap=tk.WORD)
scr.grid(row = 3, column = 1, columnspan = 2, padx=5, pady=5)

left = tkinter.StringVar()
now_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
left_num = 1000
string = '每月各项服务免费仅千次，请酌情使用'
left.set(string)
left_frame = ttk.LabelFrame(main_window, text = '剩余次数')
left_frame.grid(row = 4, column = 0, padx = 6, pady=5, sticky = W)
ttk.Entry(left_frame, width = 29, textvariable = left, state = DISABLED).grid(row = 4, column = 0, padx=5, pady=5)


out_format = tkinter.IntVar()
out_frame = ttk.LabelFrame(main_window, text = '输出格式')
out_frame.grid(row = 4, column = 1, rowspan = 2, pady=5, sticky = W+N)
tkinter.Radiobutton(out_frame, variable = out_format, value ='1', text='TXT').grid(row = 5, column = 1, pady=3)
tkinter.Radiobutton(out_frame, variable = out_format, value ='2', text='DOC').grid(row = 5, column = 2, pady=3)

using_key = tkinter.StringVar()
key_frame = ttk.LabelFrame(main_window, text = 'Using Key')
key_frame.grid(row = 4, column = 2, padx=5, pady=5)
ttk.Entry(key_frame, width = 10, textvariable = using_key).grid(row = 5, column = 3, padx=5, pady=5)
ttk.Button(key_frame, text = "了解", width = 8, command = UsingKeyTip).grid(row = 5, column = 4)

vehicle_license = StringVar()
business_card = StringVar()
vehicle_license.set("行驶证页面选择")
business_card.set("名片图像修正")
extra_frame = ttk.LabelFrame(main_window, text = '额外参数')
extra_frame.grid(row = 6, column = 0, columnspan = 4, pady=5, sticky = N)
ttk.Button(extra_frame, width = 30, text = "身份证识别相关选项", command = TK_IDCard).grid(row = 6, column = 0, padx=10, pady=5)
ttk.OptionMenu(extra_frame, vehicle_license, "行驶证页面选择", "行驶证主页正面", "行驶证副页正面").grid(row = 6, column = 1, padx=10, pady=5)
ttk.OptionMenu(extra_frame, business_card, "名片图像修正", "是", "否").grid(row = 6, column = 2, padx=10, pady=5)

tk.Button(main_window, width = 30, height = 2, relief = GROOVE, text = "开始", fg="red", command = Start).grid(row = 7, column = 0)
other_frame = ttk.LabelFrame(main_window, text = '提示&关于')
other_frame.grid(row = 7, column = 1, columnspan = 2, padx=5, pady=5)
ttk.Button(other_frame, width = 15, text = "提示", command = Tip).grid(row = 7, column = 1, padx=10, pady=5)
ttk.Button(other_frame, width = 15, text = "关于", command = About).grid(row = 7, column = 2, padx=10, pady=5)

main_window.mainloop()

'''
if __name__ == '__main__':
    while True:
        data = input("input data:")
        socket_client_send_data(data)
'''
